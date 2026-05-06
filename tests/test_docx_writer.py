import io

from docx import Document

from src.utils.docx_writer import markdown_to_docx_bytes


def _open(b: bytes) -> Document:
    return Document(io.BytesIO(b))


def test_returns_valid_docx_zip():
    b = markdown_to_docx_bytes("## H\nbody")
    assert b[:2] == b"PK"  # zip magic bytes (.docx is a zip)


def test_title_renders_as_top_heading():
    b = markdown_to_docx_bytes("body", title="My Title")
    doc = _open(b)
    # python-docx maps level=0 to the "Title" style, not "Heading 0"
    titles = [p.text for p in doc.paragraphs if p.style.name == "Title"]
    assert "My Title" in titles


def test_h2_h3_rendered_with_correct_levels():
    b = markdown_to_docx_bytes("## Big\n### Small")
    doc = _open(b)
    h2 = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
    h3 = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
    assert any(p.text == "Big" for p in h2)
    assert any(p.text == "Small" for p in h3)


def test_bullets_use_list_bullet_style():
    b = markdown_to_docx_bytes("- alpha\n- beta")
    doc = _open(b)
    bullets = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert "alpha" in bullets
    assert "beta" in bullets


def test_bold_runs_are_bold():
    b = markdown_to_docx_bytes("**important** plain")
    doc = _open(b)
    for p in doc.paragraphs:
        if "important" in p.text:
            for run in p.runs:
                if "important" in run.text:
                    assert run.bold is True
                    return
    raise AssertionError("expected a bold run containing 'important'")


def test_thai_content_preserved():
    b = markdown_to_docx_bytes("## หัวข้อภาษาไทย\nเนื้อหา **ตัวหนา**", title="ทดสอบ")
    doc = _open(b)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "หัวข้อภาษาไทย" in full_text
    assert "เนื้อหา" in full_text
    assert "ทดสอบ" in full_text
