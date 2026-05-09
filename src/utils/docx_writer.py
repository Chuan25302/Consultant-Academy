"""
Naive Markdown -> DOCX bytes for the Knowledge Base archive.
Handles: ## h2, ### h3, **bold**, - bullets, > blockquotes, plain paragraphs.
When `pillar` is supplied, the cover page and headings adopt the pillar's
brand color so KB articles are visually consistent with the daily emails.
"""
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Hex color per pillar — mirrors PILLAR_CONFIG in designer_agent.py.
# Duplicated here (rather than imported) to keep docx_writer free of the
# email-rendering dependency chain (premailer, cssutils, markdown).
PILLAR_COLOR = {
    "TECHNICAL":      RGBColor(0x0D, 0x47, 0xA1),
    "INDUSTRY":       RGBColor(0xE6, 0x51, 0x00),
    "FRAMEWORK":      RGBColor(0x1B, 0x5E, 0x20),
    "SOFTSKILL":      RGBColor(0x4A, 0x14, 0x8C),
    "COMPLIANCE":     RGBColor(0xB7, 0x1C, 0x1C),
    "SUSTAINABILITY": RGBColor(0x2E, 0x7D, 0x32),
    "RECAP":          RGBColor(0x37, 0x47, 0x4F),
}
DEFAULT_HEADING_COLOR = RGBColor(0x21, 0x21, 0x21)


def markdown_to_docx_bytes(md: str, title: str = "",
                           pillar: str | None = None,
                           subtitle: str | None = None,
                           date: str | None = None) -> bytes:
    doc = Document()
    color = PILLAR_COLOR.get((pillar or "").upper(), DEFAULT_HEADING_COLOR)

    # Cover block (when caller supplies a title): pillar tag + topic +
    # optional subtitle/date, all centered. python-docx doesn't support
    # real cover pages so we use prominent runs instead.
    if title:
        if pillar:
            tag = doc.add_paragraph()
            tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = tag.add_run(pillar.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = color

        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = color

        if subtitle:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = sub.add_run(subtitle)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        if date:
            d = doc.add_paragraph()
            d.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = d.add_run(date)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        doc.add_paragraph("")  # spacer

    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = color
        elif line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            for run in h.runs:
                run.font.color.rgb = color
        elif line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = color
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            _add_runs(p, line[2:])
        elif line.startswith(("- ", "* ", "+ ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:])
        elif line.lstrip().startswith(("* ", "- ")) and line.startswith((" ", "\t")):
            # Indented bullet — strip leading whitespace + marker
            stripped = line.lstrip()
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs(paragraph, text: str):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
